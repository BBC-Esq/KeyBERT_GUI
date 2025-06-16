import os
from pathlib import Path
from PySide6.QtCore import QThread, Signal

class TextExtractorWorker(QThread):
   text_extracted = Signal(str)
   extraction_error = Signal(str)
   finished = Signal()

   def __init__(self, file_path):
       super().__init__()
       self.file_path = file_path

   def run(self):
       try:
           text = self.extract_text_from_file(self.file_path)
           self.text_extracted.emit(text)
       except Exception as e:
           self.extraction_error.emit(str(e))
       finally:
           self.finished.emit()

   def extract_text_from_file(self, file_path):
       file_extension = Path(file_path).suffix.lower()

       if file_extension == '.txt':
           return self._extract_from_txt(file_path)
       elif file_extension == '.pdf':
           return self._extract_from_pdf(file_path)
       elif file_extension in ['.doc', '.docx']:
           return self._extract_from_word(file_path)
       else:
           raise ValueError(f"Unsupported file format: {file_extension}")

   def _extract_from_txt(self, file_path):
       try:
           with open(file_path, 'r', encoding='utf-8') as file:
               return file.read()
       except UnicodeDecodeError:
           encodings = ['latin-1', 'cp1252', 'iso-8859-1']
           for encoding in encodings:
               try:
                   with open(file_path, 'r', encoding=encoding) as file:
                       return file.read()
               except UnicodeDecodeError:
                   continue
           raise ValueError("Could not decode the text file with common encodings")

   def _extract_from_pdf(self, file_path):
       try:
           import fitz
       except ImportError:
           raise ImportError("PyMuPDF is required for PDF extraction. Install with: pip install PyMuPDF")
       
       text = ""
       try:
           doc = fitz.open(file_path)
           for page in doc:
               text += page.get_text() + "\n"
           doc.close()
       except Exception as e:
           raise Exception(f"Failed to extract PDF with PyMuPDF: {str(e)}")
       
       if not text.strip():
           raise ValueError("No text could be extracted from the PDF file")
       
       return text.strip()

   def _extract_from_word(self, file_path):
       file_extension = Path(file_path).suffix.lower()
       
       if file_extension == '.docx':
           return self._extract_from_docx(file_path)
       elif file_extension == '.doc':
           return self._extract_from_doc(file_path)
       else:
           raise ValueError(f"Unsupported Word format: {file_extension}")

   def _extract_from_docx(self, file_path):
       try:
           import docx
       except ImportError:
           raise ImportError("python-docx is required for DOCX extraction. Install with: pip install python-docx")
       
       try:
           doc = docx.Document(file_path)
           text = ""
           for paragraph in doc.paragraphs:
               text += paragraph.text + "\n"
           
           for table in doc.tables:
               for row in table.rows:
                   for cell in row.cells:
                       text += cell.text + " "
                   text += "\n"
           
           if not text.strip():
               raise ValueError("No text could be extracted from the DOCX file")
           
           return text.strip()
       except Exception as e:
           raise Exception(f"Failed to extract text from DOCX file: {str(e)}")

   def _extract_from_doc(self, file_path):
       try:
           import textract
       except ImportError:
           raise ImportError("textract is required for DOC extraction. Install with: pip install textract")
       
       try:
           text = textract.process(file_path).decode('utf-8')
           if not text.strip():
               raise ValueError("No text could be extracted from the DOC file")
           return text.strip()
       except Exception as e:
           try:
               import docx2txt
               text = docx2txt.process(file_path)
               if not text.strip():
                   raise ValueError("No text could be extracted from the DOC file")
               return text.strip()
           except ImportError:
               raise Exception(f"Failed to extract text from DOC file. Consider installing docx2txt: pip install docx2txt. Original error: {str(e)}")
           except Exception as fallback_error:
               raise Exception(f"Failed to extract DOC with both textract and docx2txt: {str(e)}, {str(fallback_error)}")


class TextExtractor:
   
   @staticmethod
   def extract_text(file_path):
       worker = TextExtractorWorker(file_path)
       return worker.extract_text_from_file(file_path)

   @staticmethod
   def get_supported_formats():
       return ['.txt', '.pdf', '.doc', '.docx']

   @staticmethod
   def is_supported_format(file_path):
       extension = Path(file_path).suffix.lower()
       return extension in TextExtractor.get_supported_formats()