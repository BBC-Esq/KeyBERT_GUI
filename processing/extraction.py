from __future__ import annotations

from pathlib import Path

from PySide6.QtCore import QThread, Signal

class TextExtractorWorker(QThread):
    """Background thread that extracts text from TXT, PDF, or DOCX."""

    text_extracted = Signal(str)
    extraction_error = Signal(str)
    finished = Signal()

    def __init__(self, file_path: str | Path) -> None:
        super().__init__()
        self.file_path = Path(file_path)

    # ------------------------------
    def run(self) -> None:
        try:
            text = self._extract_from_file(self.file_path)
            self.text_extracted.emit(text)
        except Exception as exc:
            self.extraction_error.emit(str(exc))
        finally:
            self.finished.emit()

    # Static helpers for each format
    def _extract_from_file(self, path: Path) -> str:
        ext = path.suffix.lower()
        if ext == ".txt":
            return self._extract_txt(path)
        if ext == ".pdf":
            return self._extract_pdf(path)
        if ext == ".docx":
            return self._extract_docx(path)
        raise ValueError(f"Unsupported file type: {ext}")

    def _extract_txt(self, path: Path) -> str:
        encodings = ["utf-8", "latin-1", "cp1252", "iso-8859-1"]
        for enc in encodings:
            try:
                return path.read_text(encoding=enc)
            except UnicodeDecodeError:
                continue
        raise ValueError("Could not decode TXT file with common encodings")

    # PDF
    def _extract_pdf(self, path: Path) -> str:
        try:
            import fitz
        except ImportError:
            raise ImportError("Install PyMuPDF:  pip install PyMuPDF")
        try:
            with fitz.open(path) as doc:
                text = "\n".join(page.get_text("text") for page in doc)
        except Exception as exc:
            raise RuntimeError(f"Failed to read PDF: {exc}") from exc
        if not text.strip():
            raise ValueError("No extractable text in PDF")
        return text.strip()

    # DOCX
    def _extract_docx(self, path: Path) -> str:
        try:
            import docx
        except ImportError:
            raise ImportError("Install python-docx:  pip install python-docx")
        doc = docx.Document(path)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        text = "\n".join(parts)
        if not text.strip():
            raise ValueError("No extractable text in DOCX")
        return text.strip()
