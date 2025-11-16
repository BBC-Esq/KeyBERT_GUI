from __future__ import annotations

from pathlib import Path

from PySide6.QtCore import QThread, Signal
from keybert import KeyBERT

from .extraction import TextExtractorWorker
from .output_formatter import BatchResultFormatter


class BatchProcessorWorker(QThread):

    progress_update = Signal(int, int, str)
    file_completed = Signal(str, bool)
    batch_completed = Signal(str, dict)
    error_occurred = Signal(str)

    def __init__(
        self,
        file_paths: list[Path],
        output_path: Path,
        keybert_model: KeyBERT,
        extraction_params: dict
    ) -> None:
        super().__init__()
        self.file_paths = file_paths
        self.output_path = output_path
        self.keybert_model = keybert_model
        self.extraction_params = extraction_params
        self.formatter = BatchResultFormatter()
        self._cancelled = False
    
    def cancel(self) -> None:
        self._cancelled = True

    def run(self) -> None:
        try:
            self.formatter.set_processing_parameters(self.extraction_params)
            total_files = len(self.file_paths)

            for i, file_path in enumerate(self.file_paths):
                if self._cancelled:
                    break

                self.progress_update.emit(i + 1, total_files, file_path.name)

                extracted_text = None
                keywords = None
                error = None

                try:
                    extracted_text = self._extract_text_from_file(file_path)
                    
                    if self._cancelled:
                        break
                    
                    if extracted_text and extracted_text.strip():
                        keywords = self._extract_keywords(extracted_text)
                    else:
                        error = "No text could be extracted from file"

                except Exception as e:
                    error = str(e)

                if self._cancelled:
                    break

                self.formatter.add_result(
                    file_path=file_path,
                    extracted_text=extracted_text,
                    keywords=keywords,
                    error=error
                )

                self.file_completed.emit(str(file_path), error is None)

            if not self._cancelled:
                self.formatter.save_to_file(self.output_path)
                summary_stats = self.formatter.get_summary_stats()
                self.batch_completed.emit(str(self.output_path), summary_stats)

        except Exception as e:
            self.error_occurred.emit(f"Batch processing failed: {e}")

    def _extract_text_from_file(self, file_path: Path) -> str:
        ext = file_path.suffix.lower()

        if ext == ".txt":
            return self._extract_txt(file_path)
        elif ext == ".pdf":
            return self._extract_pdf(file_path)
        elif ext == ".docx":
            return self._extract_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    def _extract_txt(self, path: Path) -> str:
        encodings = ["utf-8", "latin-1", "cp1252", "iso-8859-1"]
        for enc in encodings:
            if self._cancelled:
                raise InterruptedError("Extraction cancelled")
            try:
                return path.read_text(encoding=enc)
            except UnicodeDecodeError:
                continue
        raise ValueError("Could not decode TXT file with common encodings")

    def _extract_pdf(self, path: Path) -> str:
        try:
            import fitz
        except ImportError:
            raise ImportError("Install PyMuPDF: pip install PyMuPDF")

        try:
            with fitz.open(path) as doc:
                text_parts = []
                for page in doc:
                    if self._cancelled:
                        raise InterruptedError("Extraction cancelled")
                    text_parts.append(page.get_text("text"))
                text = "\n".join(text_parts)
        except Exception as exc:
            raise RuntimeError(f"Failed to read PDF: {exc}") from exc

        if not text.strip():
            raise ValueError("No extractable text in PDF")
        return text.strip()

    def _extract_docx(self, path: Path) -> str:
        try:
            import docx
        except ImportError:
            raise ImportError("Install python-docx: pip install python-docx")

        if self._cancelled:
            raise InterruptedError("Extraction cancelled")

        doc = docx.Document(path)
        parts = [p.text for p in doc.paragraphs]
        
        if self._cancelled:
            raise InterruptedError("Extraction cancelled")
        
        for table in doc.tables:
            if self._cancelled:
                raise InterruptedError("Extraction cancelled")
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)

        text = "\n".join(parts)
        if not text.strip():
            raise ValueError("No extractable text in DOCX")
        return text.strip()

    def _extract_keywords(self, text: str) -> list[tuple[str, float]]:
        if self._cancelled:
            raise InterruptedError("Extraction cancelled")
        
        params = self.extraction_params.copy()

        keybert_params = {
            k: v for k, v in params.items() 
            if k in ['keyphrase_ngram_range', 'stop_words', 'top_n', 'use_maxsum', 
                     'use_mmr', 'diversity', 'nr_candidates']
        }

        return self.keybert_model.extract_keywords(text, **keybert_params)